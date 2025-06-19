import os
import sys
import tkinter as tk
from tkinter import END, filedialog, messagebox
from tkinter import ttk
import speech_recognition as sr
from googletrans import Translator
from gtts import gTTS
from pydub import AudioSegment
from docx import Document
import pythoncom
import pygame

# Initialize COM and Pygame Mixer
pythoncom.CoInitialize()
pygame.mixer.init()

# Recognizer
r = sr.Recognizer()

# Create main window
main = tk.Tk()
main.title("Advanced Multilingual Translation Platform")
main.geometry("1000x750")
main.config(bg="#1E1E1E")
main.resizable(False, False)

# Languages
lt = ["English", "Hindi", "Tamil", "Gujarati", "Marathi", "Telugu", "Bengali", "Kannada", "Malayalam",
      "French", "Spanish", "German", "Chinese", "Japanese", "Russian", "Italian", "Portuguese", "Dutch",
      "Turkish", "Arabic", "Korean", "Swedish", "Norwegian", "Danish", "Polish", "Czech", "Urdu"]
code = ["en", "hi", "ta", "gu", "mr", "te", "bn", "kn", "ml", "fr", "es", "de", "zh-CN", "ja", "ru",
        "it", "pt", "nl", "tr", "ar", "ko", "sv", "no", "da", "pl", "cs", "ur"]

v1 = tk.StringVar(main)
v1.set(lt[0])  # Input language default: English

v2 = tk.StringVar(main)
v2.set(lt[1])  # Output language default: Hindi

# UI Setup
ttk.Label(main, text="Advanced Multilingual Translation Platform", font=("Arial", 16, "bold"),
          background="#1E1E1E", foreground="white").pack(pady=10)

frame1 = tk.Frame(main, bg="#333333", padx=10, pady=10)
frame1.pack(pady=10, fill="x")

ttk.Label(frame1, text="Input Text:", background="#333333", foreground="white").grid(row=0, column=0, padx=10, pady=5)
input_text = tk.Text(frame1, height=10, width=50, font=("Arial", 12))
input_text.grid(row=1, column=0, padx=10, pady=5)

ttk.Label(frame1, text="Translated Text:", background="#333333", foreground="white").grid(row=0, column=1, padx=10, pady=5)
output_text = tk.Text(frame1, height=10, width=50, font=("Arial", 12))
output_text.grid(row=1, column=1, padx=10, pady=5)

def create_button(parent, text, command, row, column):
    btn = tk.Button(parent, text=text, command=command,
                    font=("Arial", 12, "bold"), bg="#3498db", fg="white",
                    activebackground="#2980b9", activeforeground="white",
                    padx=15, pady=8, bd=0, relief="ridge", cursor="hand2")
    btn.grid(row=row, column=column, padx=10, pady=10, sticky="ew")
    return btn

def upload_text_file():
    file_path = filedialog.askopenfilename(filetypes=[("Text/DOCX Files", "*.txt *.docx"), ("All Files", "*.*")])
    if not file_path:
        return
    input_text.delete("1.0", END)
    try:
        if file_path.lower().endswith(".txt"):
            with open(file_path, "r", encoding="utf-8", errors="replace") as file:
                input_text.insert(END, file.read())
        elif file_path.lower().endswith(".docx"):
            doc = Document(file_path)
            content = [p.text for p in doc.paragraphs if p.text.strip()]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            content.append(cell.text)
            input_text.insert(END, "\n".join(content))
        else:
            messagebox.showerror("Error", "Unsupported file format")
    except Exception as e:
        messagebox.showerror("Error", str(e))

def upload_audio_file():
    file_path = filedialog.askopenfilename(filetypes=[("Audio Files", "*.wav *.mp3 *.ogg *.flac")])
    if not file_path:
        return
    input_text.delete("1.0", END)
    try:
        lang = code[lt.index(v1.get())]
        temp_file = "temp_converted.wav"
        if file_path.lower().endswith(".mp3"):
            AudioSegment.from_mp3(file_path).export(temp_file, format="wav")
            file_path = temp_file
        elif file_path.lower().endswith((".ogg", ".flac")):
            AudioSegment.from_file(file_path).export(temp_file, format="wav")
            file_path = temp_file
        with sr.AudioFile(file_path) as source:
            audio = r.record(source)
            text = r.recognize_google(audio, language=lang)
            input_text.insert(END, text)
        if os.path.exists(temp_file):
            os.remove(temp_file)
    except Exception as e:
        if 'temp_file' in locals() and os.path.exists(temp_file):
            os.remove(temp_file)
        messagebox.showerror("Audio Error", str(e))

def listen_and_translate():
    input_text.delete("1.0", END)
    output_text.delete("1.0", END)
    input_lang = code[lt.index(v1.get())]
    output_lang = code[lt.index(v2.get())]
    try:
        with sr.Microphone() as source:
            messagebox.showinfo("Speak Now", "Listening... Speak clearly into the microphone")
            r.adjust_for_ambient_noise(source, duration=1)
            audio = r.listen(source, timeout=10, phrase_time_limit=15)
            spoken_text = r.recognize_google(audio, language=input_lang)
            input_text.insert(END, spoken_text)
            translator = Translator()
            translated = translator.translate(spoken_text, dest=output_lang)
            output_text.insert(END, translated.text)
    except Exception as e:
        messagebox.showerror("Speech Error", str(e))

def translate():
    output_text.delete("1.0", END)
    text = input_text.get("1.0", END).strip()
    if not text:
        messagebox.showwarning("Warning", "Please enter text to translate")
        return
    lang = code[lt.index(v2.get())]
    try:
        translator = Translator()
        max_chunk_size = 5000
        chunks = [text[i:i+max_chunk_size] for i in range(0, len(text), max_chunk_size)]
        translated_text = [translator.translate(chunk, dest=lang).text for chunk in chunks]
        output_text.insert(END, " ".join(translated_text))
    except Exception as e:
        messagebox.showerror("Translation Error", str(e))

def speak():
    text = output_text.get("1.0", END).strip()
    if not text:
        messagebox.showwarning("Warning", "No translated text to speak")
        return
    lang = code[lt.index(v2.get())]
    try:
        tts = gTTS(text=text, lang=lang, slow=False)
        if not os.path.exists("temp"):
            os.makedirs("temp")
        temp_file = os.path.join("temp", "output.mp3")
        tts.save(temp_file)
        pygame.mixer.music.load(temp_file)
        pygame.mixer.music.play()
    except Exception as e:
        messagebox.showerror("Speech Error", str(e))

def stop_speaking():
    if pygame.mixer.music.get_busy():
        pygame.mixer.music.stop()

def save_translated_text():
    text = output_text.get("1.0", END).strip()
    if not text:
        messagebox.showwarning("Warning", "No translated text to save")
        return
    file_path = filedialog.asksaveasfilename(defaultextension=".txt",
                                             filetypes=[("Text files", "*.txt"), ("Word Documents", "*.docx")])
    if not file_path:
        return
    try:
        if file_path.endswith(".docx"):
            doc = Document()
            for line in text.splitlines():
                doc.add_paragraph(line)
            doc.save(file_path)
        else:
            with open(file_path, "w", encoding="utf-8") as f:
                f.write(text)
        messagebox.showinfo("Saved", "File saved successfully")
    except Exception as e:
        messagebox.showerror("Save Error", str(e))

def save_voice():
    text = output_text.get("1.0", END).strip()
    if not text:
        messagebox.showwarning("Warning", "No text to convert")
        return
    lang = code[lt.index(v2.get())]
    file_path = filedialog.asksaveasfilename(defaultextension=".mp3",
                                             filetypes=[("MP3 files", "*.mp3"), ("WAV files", "*.wav")])
    if not file_path:
        return
    try:
        gTTS(text=text, lang=lang).save(file_path)
        messagebox.showinfo("Success", "Audio file saved")
    except Exception as e:
        messagebox.showerror("Audio Save Error", str(e))

# Buttons
button_frame = tk.Frame(main, bg="#1E1E1E")
button_frame.pack(pady=10)
for i in range(3):
    button_frame.grid_columnconfigure(i, weight=1)

buttons = [
    ("\U0001F4C4 Upload File (TXT/DOCX)", upload_text_file),
    ("\U0001F3A4 Upload Audio File", upload_audio_file),
    ("\U0001F3A7 Listen and Translate", listen_and_translate),
    ("\U0001F4DD Translate", translate),
    ("\U0001F50A Speak", speak),
    ("\U0001F507 Stop Speaking", stop_speaking),
    ("\U0001F4BE Save Text", save_translated_text),
    ("\U0001F3B5 Save Audio", save_voice),
]

for index, (text, command) in enumerate(buttons):
    create_button(button_frame, text, command, index // 3, index % 3)

# Language Selectors
lang_frame = tk.Frame(main, bg="#1E1E1E")
lang_frame.pack(pady=10)

ttk.Label(lang_frame, text="Select Input Audio Language:", background="#1E1E1E", foreground="white").pack(side="left")
ttk.Combobox(lang_frame, textvariable=v1, values=lt, state="readonly", width=20).pack(side="left", padx=10)

ttk.Label(lang_frame, text="Select Output Language:", background="#1E1E1E", foreground="white").pack(side="left")
ttk.Combobox(lang_frame, textvariable=v2, values=lt, state="readonly", width=20).pack(side="left", padx=10)

# Run App
if __name__ == "__main__":
    try:
        main.mainloop()
    finally:
        pygame.mixer.quit()
        if 'pythoncom' in globals():
            pythoncom.CoUninitialize()